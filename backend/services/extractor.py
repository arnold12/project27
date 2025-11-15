"""
Text extraction service for PDF and DOCX documents.
"""
import os
from typing import Optional
from pathlib import Path

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from backend.core.utils import get_file_extension, clean_text


class TextExtractor:
    """Service for extracting text from PDF and DOCX documents."""
    
    def __init__(self):
        """Initialize the text extractor."""
        self.preferred_pdf_library = "pymupdf" if PYMUPDF_AVAILABLE else "pdfplumber"
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text, or None if extraction fails
        """
        file_ext = get_file_extension(file_path)
        
        if file_ext == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_ext == "docx":
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """
        Extract text from PDF using PyMuPDF or pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text, or None if extraction fails
        """
        if self.preferred_pdf_library == "pymupdf" and PYMUPDF_AVAILABLE:
            return self._extract_with_pymupdf(file_path)
        elif PDFPLUMBER_AVAILABLE:
            return self._extract_with_pdfplumber(file_path)
        else:
            raise RuntimeError(
                "No PDF extraction library available. "
                "Please install PyMuPDF (fitz) or pdfplumber."
            )
    
    def _extract_with_pymupdf(self, file_path: str) -> str:
        """
        Extract text using PyMuPDF (fitz).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        full_text = "\n".join(text_parts)
        return clean_text(full_text)
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """
        Extract text using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        full_text = "\n".join(text_parts)
        return clean_text(full_text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "python-docx is not installed. "
                "Please install it to support DOCX files."
            )
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        return clean_text(full_text)
    
    def extract_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with metadata
        """
        metadata = {
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "file_extension": get_file_extension(file_path),
        }
        
        # Try to extract PDF metadata
        if metadata["file_extension"] == "pdf" and PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                pdf_metadata = doc.metadata
                metadata.update({
                    "title": pdf_metadata.get("title", ""),
                    "author": pdf_metadata.get("author", ""),
                    "subject": pdf_metadata.get("subject", ""),
                    "creator": pdf_metadata.get("creator", ""),
                    "producer": pdf_metadata.get("producer", ""),
                    "creation_date": pdf_metadata.get("creationDate", ""),
                    "modification_date": pdf_metadata.get("modDate", ""),
                    "page_count": len(doc),
                })
                doc.close()
            except Exception as e:
                print(f"Error extracting PDF metadata: {e}")
        
        return metadata


# Global extractor instance
text_extractor = TextExtractor()

